from view.generate_request import GenerateRequestModel
from controller.gpt_exam_generator import GPT_ExamGenerator
from controller.nlp import NLP
from docx import Document
import re
import os

class _ReadCourseExams:
    def __init__(self):
        self.root = 'question_bank'
        return
    def set_course(self, course:str):
        self.course = course
        return
    def check_path(self)->bool:
        print('Check 1')
        #There is no Question bank
        if not os.path.exists(f'./{self.root}'):
            print('Check 3')
            return False
        #There is Question bank, but there is not course (092*-***)
        if not os.path.exists(f'./{self.root}/{self.course}'):
            return False
        #There is question bank for the given course
        print('Check 2')
        return True
    def __readFile(self,path:str)->list:
        exam = Document(path)
        questions = []
        for p in exam.paragraphs:
            if not p.style.name.startswith('Heading'):
                question = p.text.strip()
                if question:
                    match = re.search(r'(\d+)\.\s*(.*)', question)
                    if match:
                        questions.append(match.group(2).strip())
        return questions
    def getAll(self)->list:
        all_qs = []
        for fname in os.listdir(f"./{self.root}/{self.course}"):
            if fname.endswith(".docx"):
                path = os.path.join(f'./{self.root}/{self.course}',fname)
                all_qs.extend(self.__readFile(path))
        return all_qs        
    
class _WriteCourseExams:
    def __init__(self):
        self.root = 'question_bank'
        return
    
    def set_course(self, course:str)->None:
        if course is None:
            raise Exception("Course Should Not Be Empty!")
        self.course = course
        return
    
    def __nextFileName(self):
        pattern = re.compile(rf'{self.course}-e(\d+)\.docx')
        last = 0
        #In case root folder of exams (course) does not exist
        if not os.path.exists(f"./{self.root}"):
            os.makedirs(f"./{self.root}/{self.course}")
            return (os.path.join(f"./{self.root}/{self.course}", f"{self.course}-e{1}.docx"), f"{self.course}-e{1}.docx")
        
        #In case course folder (092*-***) does not exist
        if not os.path.exists(f"./{self.root}/{self.course}"):
            os.mkdir(f'./{self.root}/{self.course}')
            return (os.path.join(f"./{self.root}/{self.course}", f"{self.course}-e{1}.docx"), f"{self.course}-e{1}.docx")
        
        for fname in os.listdir(f"./{self.root}/{self.course}"):
            match = pattern.match(fname)
            if match :
                last = max(last, int(match.group(1)))
        new_file = last + 1
        print((os.path.join(f"./{self.root}/{self.course}", f"{self.course}-e{new_file}.docx")))
        print( f"{self.course}-e{new_file}.docx")
        return ((os.path.join(f"./{self.root}/{self.course}", f"{self.course}-e{new_file}.docx")), f"{self.course}-e{new_file}.docx")
    
    def split_Qs(self, exam:str)->tuple:
        # questions_text = re.sub(r'\s+', ' ', exam)
        MCQ = r'###\s*Multiple-Choice Questions(.*?)(?=###|$)' #MCQ respnse pattern
        TF = r'###\s*True/False Questions(.*?)(?=###|$)' #True/False Questions response pattern
        SA = r'###\s*True/False Questions(.*?)(?=###|$)' #Short Answers Questions response pattern
        mcsqs = re.findall(MCQ, exam, re.DOTALL)
        tfqs = re.findall(TF, exam, re.DOTALL)
        saqs = re.findall(SA, exam, re.DOTALL)
        return (mcsqs, tfqs, saqs) # return questions grouped by their type
    
    def writeExam(self, exam:str):
        new_document = Document()
        sections = exam.split("##")
        for section in sections[1:]:  # Skip the first split which will be empty
            print("Hi Nabaa")
            title, *content = section.split('\n', 1)
            content = content[0] if content else ''
            
            # Add section heading
            new_document.add_heading(title.strip(), level=1)
            
            # Process each question
            questions = content.strip().split('\n\n')
            for question in questions:
                lines = question.split('\n')
                question_text = lines[0].strip()
                new_document.add_paragraph(question_text, style='ListParagraph')

                # Handle multiple-choice separately to format options
                if 'Multiple-choice' in title:
                    for line in lines[1:]:
                        if line.strip().startswith('- **Answer:'):
                            # Add answer text
                            answer_text = line.strip()[2:]  # Remove leading '- '
                            new_document.add_paragraph(answer_text, style='IntenseQuote')
                        else:
                            # Add options
                            new_document.add_paragraph(line.strip(), style='BodyText')
                else:
                    for line in lines[1:]:
                        new_document.add_paragraph(line.strip(), style='BodyText')
        
        (path, name) = self.__nextFileName()
        print(path)
        new_document.save(path)
        return (path,name)
    
    def replace_question(self, old_q:str, new_q:str, docpath:str):
        doc = Document(docpath)
        for p in doc.paragraphs:
            if old_q in p.text:
                p.clear()
                p.add_run(new_q)
        doc.save()
        return
    

class CourseExams:
    def __init__(self, course:str) -> None:
        self.writer = _WriteCourseExams()
        self.reader = _ReadCourseExams()
        self.writer.set_course(course=course)
        self.reader.set_course(course=course)
        self.nlp = NLP()
        return
    
    def write(self,exam:str):
        return self.writer.writeExam(exam=exam)
    
    def read(self):
        if self.reader.check_path():
            self.questions = self.reader.getAll()
            print("Old Exams, Found!")
        else:
            print("No Old Exams, You Can write!")
        return
    def __split(self,newQs)->list:
        pattern = r'\*\*Question Type: (.*?)\*\*\n(.*?)\n\*\*Answer:'
        matches = [match.strip() for match in re.findall(pattern=pattern, string= newQs , flags=re.DOTALL)]
        newOnes = []
        for i in matches:
            match = re.search(r'\*\*Question:\*\*\s*(.*)', i)
            type = re.search(r'\*\*Question Type:\s*(.*?)\*\*', i)
            if match:
                newOnes.append((match.group(1).strip(), type.group(1).strip()))
        return newOnes
    def check_similarities(self, newQs:str, threshold=0.75):
        pairs = []
        newQsList = self.__split(newQs=newQs)
        similarities = self.nlp.isSimilar(newQ=newQsList,oldQ=self.questions)
        print(f'similarities Matrix:\n{similarities}')
        if len(similarities) != 0:
            for i, (n,type) in enumerate(newQsList):
                for j,o in enumerate(self.questions):
                    score = similarities[i][j].item()
                    if score >= threshold:
                        paraphrased = self.paraphrase(n)
                        paraphrased_similarity_score = self.nlp.isSimilar([paraphrased], [o])
                        if paraphrased_similarity_score[0][0] >= threshold:
                            pairs.append((n,o, type))
        return pairs
    def replace_similar(self, newQs:str, request:GenerateRequestModel, content:str, fpath:str):
        similar_pairs = self.check_similarities(newQs=newQs)
        for new_q, old_q, type in similar_pairs:
            difficulty = ''
            if type == "Multiple Choice":
                difficulty = request.questions.MCQ.difficulty
            elif type == "True/False":
                difficulty = request.questions.TF.difficulty
            else:
                difficulty = request.questions.ShortAnswer.difficulty
            gen_q = self.generate_new(question=new_q, difficulty=difficulty, type=type, content=content)
            self.writer.replace_question(old_q=old_q, new_q=gen_q, docpath=fpath)
        return similar_pairs
    def paraphrase(self, question:str):
        return self.nlp.PM.paraphrase(q=question)
    def generate_new(question:str, difficulty:str, type:str, content:str):
        gpt = GPT_ExamGenerator()
        prompt = f"""
                You are an AI assistant helping an instructor generate an exam. Create a new question to replace the following one:
                - Original question: {question}
                - Course content: {content}
                - Question type: {type}
                - Difficulty: {difficulty}
                - The new question must be different from the old question but maintain the same format, difficulty, and type.
                - Provide answers for multiple-choice questions, True/False questions with explanation, and short answer questions based on the given content.
                - Clearly specify the question type and answer.

                Start generating a replacement question:
            """
        return gpt._send(prompt)